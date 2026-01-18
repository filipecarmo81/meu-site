# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime

# Carregar dados
df_prot = pd.read_excel(r'C:\Users\AMH\Desktop\meu-site\Relatorio_Comparacao_v3.xlsx', sheet_name='1-Resumo Protocolos')
df_contas = pd.read_excel(r'C:\Users\AMH\Desktop\meu-site\Relatorio_Comparacao_v3.xlsx', sheet_name='2-Resumo Contas')

# Criar documento
doc = Document()

# Configurar margens
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Funcao para adicionar tabela
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

# Calcular metricas
total_prot = len(df_prot)
prot_ok = len(df_prot[df_prot['STATUS'] == 'OK'])
prot_apenas_excel = len(df_prot[df_prot['STATUS'] == 'APENAS EXCEL'])
prot_apenas_xml = len(df_prot[df_prot['STATUS'] == 'APENAS XML'])
prot_duplicados = len(df_prot[df_prot['STATUS'] == 'DUPLICADO'])

total_contas = len(df_contas)
contas_ok = len(df_contas[df_contas['STATUS'] == 'OK'])
contas_apenas_excel = len(df_contas[df_contas['STATUS'] == 'APENAS EXCEL'])
contas_apenas_xml = len(df_contas[df_contas['STATUS'] == 'APENAS XML'])
contas_duplicadas = len(df_contas[df_contas['STATUS'] == 'DUPLICADO'])

pct_prot_ok = (prot_ok / total_prot * 100)
pct_contas_ok = (contas_ok / total_contas * 100)

# Contas orfas (protocolo tem XML mas conta nao)
contas_sem_xml = df_contas[df_contas['STATUS'] == 'APENAS EXCEL']
protocolos_com_xml = set(df_prot[df_prot['NO_XML'] == 'Sim']['NR_SEQ_PROTOCOLO'].astype(str))
contas_orfas = contas_sem_xml[contas_sem_xml['NR_SEQ_PROTOCOLO'].astype(str).isin(protocolos_com_xml)]
contas_esperadas = contas_sem_xml[~contas_sem_xml['NR_SEQ_PROTOCOLO'].astype(str).isin(protocolos_com_xml)]

# ==================== CAPA ====================
title = doc.add_heading('Analise de Consistencia', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Contas Medicas: Excel (TASY) vs XML (TISS)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
date_para = doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== SUMARIO EXECUTIVO ====================
doc.add_heading('1. Sumario Executivo', level=1)

p = doc.add_paragraph()
p.add_run('Objetivo: ').bold = True
p.add_run('Verificar a consistencia entre os dados de contas medicas registrados no sistema TASY (Excel) e os arquivos XML enviados aos planos de saude.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Resultado Geral: ').bold = True
p.add_run(f'{pct_prot_ok:.1f}% dos protocolos e {pct_contas_ok:.1f}% das contas estao consistentes entre Excel e XML.')

doc.add_paragraph()

doc.add_heading('Visao Geral', level=2)
add_table(doc,
    ['Metrica', 'Protocolos', 'Contas'],
    [
        ['Total Analisado', str(total_prot), str(total_contas)],
        ['Consistentes (OK)', f'{prot_ok} ({pct_prot_ok:.1f}%)', f'{contas_ok} ({pct_contas_ok:.1f}%)'],
        ['Apenas no Excel', str(prot_apenas_excel), str(contas_apenas_excel)],
        ['Apenas no XML', str(prot_apenas_xml), str(contas_apenas_xml)],
        ['Duplicados', str(prot_duplicados), str(contas_duplicadas)],
    ]
)

doc.add_page_break()

# ==================== ANALISE DE PROTOCOLOS ====================
doc.add_heading('2. Analise de Protocolos', level=1)

doc.add_heading('2.1 Distribuicao por Status', level=2)
add_table(doc,
    ['Status', 'Quantidade', 'Percentual', 'Interpretacao'],
    [
        ['OK', str(prot_ok), f'{prot_ok/total_prot*100:.1f}%', 'Protocolo existe no Excel e no XML'],
        ['APENAS EXCEL', str(prot_apenas_excel), f'{prot_apenas_excel/total_prot*100:.1f}%', 'XML nao gerado ou ausente na pasta'],
        ['DUPLICADO', str(prot_duplicados), f'{prot_duplicados/total_prot*100:.1f}%', 'Mais de um arquivo XML'],
        ['APENAS XML', str(prot_apenas_xml), f'{prot_apenas_xml/total_prot*100:.1f}%', 'Protocolo removido do Excel'],
    ]
)

doc.add_heading('2.2 Pontos de Atencao', level=2)

if prot_apenas_excel > 0:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{prot_apenas_excel} protocolos sem XML: ').bold = True
    p.add_run('Verificar se sao protocolos recentes ainda nao fechados ou se houve falha no envio.')

if prot_duplicados > 0:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{prot_duplicados} protocolos com XML duplicado: ').bold = True
    p.add_run('Verificar se os arquivos duplicados sao identicos ou se ha divergencia de conteudo.')

if prot_apenas_xml > 0:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{prot_apenas_xml} protocolo(s) apenas no XML: ').bold = True
    p.add_run('Verificar se foram cancelados/estornados no sistema apos o envio.')

doc.add_heading('2.3 Protocolos sem XML', level=2)
lista_prot = df_prot[df_prot['STATUS'] == 'APENAS EXCEL']['NR_SEQ_PROTOCOLO'].tolist()
p = doc.add_paragraph()
p.add_run(f'Total: {len(lista_prot)} protocolos').bold = True
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Lista: ').bold = True
p.add_run(', '.join([str(x) for x in sorted(lista_prot)]))

doc.add_heading('2.4 Protocolos com XML Duplicado', level=2)
df_dup = df_prot[df_prot['STATUS'] == 'DUPLICADO'][['NR_SEQ_PROTOCOLO', 'ARQUIVOS_XML']]
if len(df_dup) > 0:
    rows = []
    for _, row in df_dup.iterrows():
        arq = str(row['ARQUIVOS_XML'])
        if len(arq) > 60:
            arq = arq[:60] + '...'
        rows.append([str(row['NR_SEQ_PROTOCOLO']), arq])
    add_table(doc, ['Protocolo', 'Arquivos XML'], rows)
else:
    doc.add_paragraph('Nenhum protocolo duplicado encontrado.')

doc.add_page_break()

# ==================== ANALISE DE CONTAS ====================
doc.add_heading('3. Analise de Contas', level=1)

doc.add_heading('3.1 Distribuicao por Status', level=2)
add_table(doc,
    ['Status', 'Quantidade', 'Percentual', 'Interpretacao'],
    [
        ['OK', str(contas_ok), f'{contas_ok/total_contas*100:.1f}%', 'Conta existe no Excel e no XML'],
        ['APENAS EXCEL', str(contas_apenas_excel), f'{contas_apenas_excel/total_contas*100:.1f}%', 'Conta nao consta em nenhum XML'],
        ['DUPLICADO', str(contas_duplicadas), f'{contas_duplicadas/total_contas*100:.1f}%', 'Conta em mais de um XML'],
        ['APENAS XML', str(contas_apenas_xml), f'{contas_apenas_xml/total_contas*100:.1f}%', 'Conta removida do Excel'],
    ]
)

doc.add_heading('3.2 Analise das Contas sem XML', level=2)
add_table(doc,
    ['Situacao', 'Quantidade', 'Observacao'],
    [
        ['Protocolo NAO tem XML', str(len(contas_esperadas)), 'Esperado - protocolo inteiro sem XML'],
        ['Protocolo TEM XML', str(len(contas_orfas)), 'INVESTIGAR - conta deveria estar no XML'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ATENCAO: ').bold = True
p.add_run(f'Existem {len(contas_orfas)} contas que pertencem a protocolos que possuem XML, porem essas contas nao estao incluidas nos arquivos XML. Isso pode indicar remocao manual ou erro na geracao.')

doc.add_heading('3.3 Pontos de Atencao', level=2)

if len(contas_orfas) > 0:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{len(contas_orfas)} contas orfas: ').bold = True
    p.add_run('Contas que deveriam estar no XML mas nao estao. Prioridade alta para investigacao.')

if contas_duplicadas > 0:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{contas_duplicadas} contas duplicadas: ').bold = True
    p.add_run('Verificar se ha cobranca em duplicidade.')

if contas_apenas_xml > 0:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{contas_apenas_xml} contas apenas no XML: ').bold = True
    p.add_run('Contas enviadas ao plano que nao constam mais no sistema.')

doc.add_page_break()

# ==================== QUESTIONAMENTOS ====================
doc.add_heading('4. Questionamentos para Investigacao', level=1)

doc.add_heading('4.1 Sobre Protocolos', level=2)
questions = [
    f'Os {prot_apenas_excel} protocolos sem XML sao recentes e ainda estao em processo de fechamento?',
    'Qual o prazo esperado entre o fechamento do protocolo e a geracao do XML?',
    f'Os {prot_duplicados} protocolos com XML duplicado foram reenviados intencionalmente? Os arquivos sao identicos?',
]
for q in questions:
    doc.add_paragraph(q, style='List Number')

doc.add_heading('4.2 Sobre Contas', level=2)
questions = [
    f'As {len(contas_orfas)} contas orfas (protocolo tem XML mas conta nao esta nele) foram removidas intencionalmente?',
    'Existe processo de auditoria para validar se todas as contas do protocolo foram incluidas no XML antes do envio?',
    f'As {contas_apenas_xml} contas que estao no XML mas nao no Excel foram estornadas? Se sim, foi enviada retificacao ao plano?',
]
for q in questions:
    doc.add_paragraph(q, style='List Number')

doc.add_heading('4.3 Sobre Processo', level=2)
questions = [
    'Existe controle de versao dos arquivos XML enviados?',
    'Ha log de alteracoes manuais nos arquivos XML antes do envio?',
    'Qual o procedimento quando uma conta precisa ser removida de um protocolo ja enviado?',
]
for q in questions:
    doc.add_paragraph(q, style='List Number')

doc.add_page_break()

# ==================== RECOMENDACOES ====================
doc.add_heading('5. Recomendacoes', level=1)

doc.add_heading('5.1 Curto Prazo', level=2)
items = [
    f'Investigar as {len(contas_orfas)} contas orfas para identificar causa raiz',
    f'Validar se os {prot_duplicados} XMLs duplicados tem conteudo identico',
    'Verificar status dos protocolos sem XML junto a area responsavel',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2 Medio Prazo', level=2)
items = [
    'Implementar validacao automatica antes do envio do XML',
    'Criar relatorio periodico de consistencia Excel vs XML',
    'Documentar processo de retificacao de XMLs enviados',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Longo Prazo', level=2)
items = [
    'Automatizar geracao de XMLs diretamente do sistema',
    'Implementar trilha de auditoria para alteracoes',
    'Integrar validacao de consistencia no fluxo de fechamento',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Salvar
doc.save(r'C:\Users\AMH\Desktop\meu-site\Analise_Executiva_Protocolos_Contas.docx')
print('Documento salvo com sucesso!')
print(r'Arquivo: C:\Users\AMH\Desktop\meu-site\Analise_Executiva_Protocolos_Contas.docx')
